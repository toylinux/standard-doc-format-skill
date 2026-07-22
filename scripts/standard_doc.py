# -*- coding: utf-8 -*-
"""标准格式 Word 文档生成器。

输入：一个纯文本文件，使用标签行标记段落类型：
    [TITLE] 文档标题（居中，方正小标宋简体 二号）
    [SIGN]  署名（居中，楷体_GB2312 三号）
    [H1]    一级标题（黑体 三号，首行缩进2字符）
    [H2]    二级标题（楷体_GB2312 三号 加粗，首行缩进2字符）
    [BODY]  正文（仿宋_GB2312 三号，首行缩进2字符，两端对齐）
    无标签的普通行按正文处理。

输出：符合以下规范的 .docx：
    页边距 上/下 2.54cm，左/右 3.18cm
    行距 固定值 30 磅
    页脚 居中 纯数字页码（小四 12pt，PAGE 域自动更新）

用法：
    python standard_doc.py <输入.txt> [输出.docx]
"""
import sys
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT_TITLE = "方正小标宋简体"
FONT_SIGN = "楷体_GB2312"
FONT_BODY = "仿宋_GB2312"
FONT_H1 = "黑体"

SIZE_ERHAO = Pt(22)    # 二号
SIZE_SANHAO = Pt(16)   # 三号
SIZE_PAGE = Pt(12)     # 小四（页码）


def set_run_font(run, font_name, size, bold=False):
    run.font.name = font_name
    run.font.size = size
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        from lxml import etree
        rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)


def add_para(doc, text, font_name, size, align, bold=False,
             indent_chars=2, first_line_indent=True):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(30)        # 固定值 30 磅
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Pt(size.pt * indent_chars) if first_line_indent else Pt(0)
    run = p.add_run(text)
    set_run_font(run, font_name, size, bold)
    return p


def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(30)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    run = p.add_run()
    set_run_font(run, FONT_BODY, SIZE_PAGE)
    fld1 = OxmlElement('w:fldChar')
    fld1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    fld2 = OxmlElement('w:fldChar')
    fld2.set(qn('w:fldCharType'), 'end')
    run._r.append(fld1)
    run._r.append(instr)
    run._r.append(fld2)


def parse_input(path):
    title = None
    sign = ""
    blocks = []
    with open(path, encoding='utf-8') as f:
        for raw in f:
            s = raw.rstrip('\n').rstrip('\r')
            if s.startswith('[TITLE]'):
                title = s[len('[TITLE]'):].strip()
            elif s.startswith('[SIGN]'):
                sign = s[len('[SIGN]'):].strip()
            elif s.startswith('[H1]'):
                blocks.append(('h1', s[len('[H1]'):].strip()))
            elif s.startswith('[H2]'):
                blocks.append(('h2', s[len('[H2]'):].strip()))
            elif s.startswith('[BODY]'):
                blocks.append(('body', s[len('[BODY]'):].strip()))
            elif s.strip():
                blocks.append(('body', s.strip()))
    return title, sign, blocks


def main():
    if len(sys.argv) < 2:
        print("用法: python standard_doc.py <输入文件.txt> [输出文件.docx]")
        sys.exit(1)
    in_path = sys.argv[1]
    out_path = sys.argv[2] if len(sys.argv) > 2 else "标准格式文档.docx"
    title, sign, blocks = parse_input(in_path)
    if not title:
        print("错误：输入文件缺少 [TITLE] 行")
        sys.exit(1)

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.18)
    sec.right_margin = Cm(3.18)

    add_para(doc, title, FONT_TITLE, SIZE_ERHAO, WD_ALIGN_PARAGRAPH.CENTER,
             indent_chars=0, first_line_indent=False)
    if sign:
        add_para(doc, sign, FONT_SIGN, SIZE_SANHAO, WD_ALIGN_PARAGRAPH.CENTER,
                 indent_chars=0, first_line_indent=False)
    for kind, text in blocks:
        if kind == 'h1':
            add_para(doc, text, FONT_H1, SIZE_SANHAO, WD_ALIGN_PARAGRAPH.LEFT)
        elif kind == 'h2':
            add_para(doc, text, FONT_SIGN, SIZE_SANHAO, WD_ALIGN_PARAGRAPH.LEFT,
                     bold=True)
        else:
            add_para(doc, text, FONT_BODY, SIZE_SANHAO, WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_page_number(sec)
    doc.save(out_path)
    print("已生成:", out_path)


if __name__ == '__main__':
    main()
